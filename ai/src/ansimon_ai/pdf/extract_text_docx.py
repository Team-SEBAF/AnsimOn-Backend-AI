import re

def _normalize_line(text: str) -> str:
    text = text.replace("\r", " ").replace("\n", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text

def _dedupe_preserve_order(lines: list[str]) -> list[str]:
    normalized_lines: list[str] = []
    seen: set[str] = set()

    for line in lines:
        if not line or line in seen:
            continue
        normalized_lines.append(line)
        seen.add(line)

    return normalized_lines

def extract_text_from_docx(docx_path: str) -> str:
    from docx import Document

    document = Document(docx_path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = _normalize_line(paragraph.text)
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [_normalize_line(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                lines.append(" | ".join(cells))

    return "\n".join(_dedupe_preserve_order(lines))